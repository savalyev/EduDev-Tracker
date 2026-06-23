# -*- coding: utf-8 -*-
"""
Генератор руководства программиста EduDev Tracker v1.0
Формат: Times New Roman 14pt, одинарный интервал, красная строка 1.25 см,
        выравнивание по ширине
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(14)
INDENT      = Cm(1.25)

# ─────────────────────────────────────────────
# Вспомогательные функции
# ─────────────────────────────────────────────

def _apply_font(run, bold=False, size=FONT_SIZE):
    run.font.name  = FONT_NAME
    run.font.size  = size
    run.font.bold  = bold
    # Для кириллицы нужно задать тему шрифта
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    FONT_NAME)
    rFonts.set(qn('w:hAnsi'),    FONT_NAME)
    rFonts.set(qn('w:cs'),       FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def _apply_para_fmt(para,
                    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    first_indent=True,
                    space_before=0,
                    space_after=0,
                    center=False):
    pf = para.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.CENTER if center else align
    pf.line_spacing_rule  = WD_LINE_SPACING.SINGLE
    pf.space_before       = Pt(space_before)
    pf.space_after        = Pt(space_after)
    pf.first_line_indent  = INDENT if first_indent else Pt(0)


def add_body(doc, text, bold=False, first_indent=True, center=False, space_before=0, space_after=0):
    """Добавить абзац тела с нужным форматированием."""
    para = doc.add_paragraph()
    _apply_para_fmt(para, first_indent=first_indent, center=center,
                    space_before=space_before, space_after=space_after)
    run = para.add_run(text)
    _apply_font(run, bold=bold)
    return para


def add_heading1(doc, text):
    """Заголовок первого уровня (раздел)."""
    para = doc.add_paragraph()
    _apply_para_fmt(para, first_indent=False, space_before=12, space_after=6)
    run = para.add_run(text)
    _apply_font(run, bold=True)
    return para


def add_heading2(doc, text):
    """Заголовок второго уровня (подраздел)."""
    para = doc.add_paragraph()
    _apply_para_fmt(para, first_indent=False, space_before=6, space_after=3)
    run = para.add_run(text)
    _apply_font(run, bold=True)
    return para


def add_table(doc, headers, rows, col_widths=None):
    """Добавить таблицу с шапкой и строками."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Ширины колонок
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # Шапка
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        _apply_font(run, bold=True)
        # серый фон шапки
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)

    # Данные
    for r_idx, row_data in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = data_row.cells[c_idx]
            para = cell.paragraphs[0]
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            run = para.add_run(str(cell_text))
            _apply_font(run)

    doc.add_paragraph()   # отступ после таблицы
    return table


def add_page_break(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run()
    run.add_break()
    br = run._r.find(qn('w:br'))
    if br is not None:
        br.set(qn('w:type'), 'page')


# ─────────────────────────────────────────────
# Создание документа
# ─────────────────────────────────────────────

doc = Document()

# Поля страницы (2 см слева/сверху/снизу, 1.5 справа — ГОСТ-ориентация)
section = doc.sections[0]
section.left_margin   = Cm(3)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ═══════════════════════════════════════════
# ТИТУЛЬНЫЙ ЛИСТ
# ═══════════════════════════════════════════

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

add_body(doc, "РУКОВОДСТВО ПРОГРАММИСТА", bold=True, first_indent=False, center=True, space_before=24)
add_body(doc, "EduDev Tracker, версия 1.0", bold=True, first_indent=False, center=True, space_before=6)

for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

add_body(doc, "Разработчик: SVVVLYXV", bold=False, first_indent=False, center=True)
add_body(doc, "Контакт: dsaval89@gmail.com", bold=False, first_indent=False, center=True)

for _ in range(8):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

add_body(doc, "2026", bold=False, first_indent=False, center=True)

add_page_break(doc)

# ═══════════════════════════════════════════
# РАЗДЕЛ 1. НАЗНАЧЕНИЕ И УСЛОВИЯ ПРИМЕНЕНИЯ
# ═══════════════════════════════════════════

add_heading1(doc, "1 НАЗНАЧЕНИЕ И УСЛОВИЯ ПРИМЕНЕНИЯ ПРИЛОЖЕНИЯ")

add_heading2(doc, "1.1 Сведения о разработчике")
add_table(doc,
    headers=["Параметр", "Значение"],
    rows=[
        ["Наименование приложения", "EduDev Tracker"],
        ["Версия",                  "1.0"],
        ["Разработчик",             "SVVVLYXV"],
        ["Контактный email",        "dsaval89@gmail.com"],
        ["Год выпуска",             "2026"],
        ["Технологии",              ".NET MAUI 10, C# 13, SQLite, XAML"],
        ["Репозиторий",             "Локально (не опубликован)"],
    ],
    col_widths=[6, 10.5],
)

add_heading2(doc, "1.2 Назначение и функции приложения")
add_body(doc, (
    "EduDev Tracker — кроссплатформенное приложение для личной продуктивности, "
    "разработанное на платформе .NET MAUI. Приложение функционирует полностью автономно: "
    "все данные хранятся локально в базе данных SQLite на устройстве пользователя. "
    "Облачная синхронизация и серверный компонент отсутствуют."
))
add_body(doc, "Приложение выполняет следующие основные функции:")
for fn in [
    "Отслеживание привычек: создание, ведение и анализ регулярных действий трёх типов (бинарных, количественных и временных), хранение истории выполнения, расчёт стриков.",
    "Управление задачами: ведение задач с приоритетами, дедлайнами, статусами, повторяющимися правилами и тремя режимами отображения (список, календарь, канбан-доска).",
    "Ведение заметок: редактор с поддержкой Markdown, полнотекстовым поиском (FTS5), категориями, вложениями файлов и версионированием содержимого.",
    "Помодоро-таймер: управление рабочими сессиями по технике «Помодоро», настраиваемые пресеты, фоновый таймер, привязка к задачам.",
    "Конвертеры данных: числовые системы, цвета (HEX/RGB/HSL), временные зоны, форматы JSON/XML, URL-кодирование.",
    "Аналитика и отчётность: агрегированная статистика по всем модулям, графики продуктивности, экспорт в XLSX и PDF.",
    "Авторизация: вход по email и паролю (BCrypt), OAuth через Google/GitHub, локальный (анонимный) режим.",
    "Управление профилем: просмотр сводной статистики, настройки приложения (язык, уведомления, звук).",
]:
    p = doc.add_paragraph()
    _apply_para_fmt(p, first_indent=False)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f"— {fn}")
    _apply_font(run)

add_heading2(doc, "1.3 Требования к аппаратному обеспечению")
add_table(doc,
    headers=["Платформа", "Параметр", "Требование"],
    rows=[
        ["Windows (Desktop)", "ОС",                     "Windows 10 Build 19041+ / Windows 11 (x64)"],
        ["Windows (Desktop)", "Оперативная память",     "4 ГБ RAM"],
        ["Windows (Desktop)", "Место на диске",         "300 МБ"],
        ["Windows (Desktop)", "Процессор",              "x64, 1 ГГц и более"],
        ["Android",           "ОС",                     "Android 8.0 (API 26) и выше"],
        ["Android",           "Оперативная память",     "2 ГБ RAM"],
        ["Android",           "Место на диске",         "150 МБ"],
        ["iOS / iPadOS",      "ОС",                     "iOS 15.0 и выше"],
        ["iOS / iPadOS",      "Место на диске",         "150 МБ"],
        ["macOS (Catalyst)",  "ОС",                     "macOS 12 (Monterey) и выше"],
        ["macOS (Catalyst)",  "Место на диске",         "200 МБ"],
    ],
    col_widths=[4, 5.5, 7],
)

add_heading2(doc, "1.4 Требования к программному обеспечению")
add_body(doc, (
    "Для сборки и запуска приложения необходима среда выполнения .NET 10 SDK (включает Runtime). "
    "Целевые фреймворки определены в файле проекта EduDev Tracker.csproj:"
))
for fw in [
    "net10.0-windows10.0.19041.0 — Windows Desktop",
    "net10.0-android            — Android",
    "net10.0-ios                — iOS",
    "net10.0-maccatalyst        — macOS (Catalyst)",
]:
    p = doc.add_paragraph()
    _apply_para_fmt(p, first_indent=False)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f"— {fw}")
    _apply_font(run)

add_body(doc, "Зависимости NuGet (основные пакеты):", space_before=4)
add_table(doc,
    headers=["Пакет", "Версия", "Назначение"],
    rows=[
        ["BCrypt.Net-Next",              "4.2.0",     "Хэширование паролей (BCrypt алгоритм)"],
        ["CommunityToolkit.Maui",        "14.1.0",    "Расширения MAUI: всплывающие окна, конверторы"],
        ["CommunityToolkit.Mvvm",        "8.4.2",     "Source-генераторы MVVM: [ObservableProperty], [RelayCommand]"],
        ["Microsoft.Maui.Controls",      "10.0.51",   "Основной фреймворк UI"],
        ["sqlite-net-pcl",               "1.9.172",   "ORM для SQLite: асинхронные CRUD-операции"],
        ["SQLiteNetExtensions",          "2.1.0",     "Реляционные связи SQLite (GetWithChildrenAsync)"],
        ["SQLiteNetExtensions.Async",    "2.1.0",     "Асинхронная версия SQLiteNetExtensions"],
        ["SQLitePCLRaw.bundle_green",    "2.1.11",    "Нативные бинарники SQLite для всех платформ"],
        ["Plugin.LocalNotification",     "14.1.0",    "Локальные push-уведомления"],
        ["Plugin.Maui.Audio",            "4.0.0",     "Воспроизведение аудио (звуки таймера)"],
        ["ClosedXML",                    "0.103.2",   "Генерация Excel-файлов (экспорт аналитики)"],
        ["Microsoft.Extensions.Logging.Debug", "10.0.0", "Логирование в DEBUG-сборках"],
    ],
    col_widths=[6.5, 2, 8],
)

add_page_break(doc)

# ═══════════════════════════════════════════
# РАЗДЕЛ 2. ХАРАКТЕРИСТИКА ПРИЛОЖЕНИЯ
# ═══════════════════════════════════════════

add_heading1(doc, "2 ХАРАКТЕРИСТИКА ПРИЛОЖЕНИЯ")

add_heading2(doc, "2.1 Архитектура и структура")
add_body(doc, (
    "EduDev Tracker построен по архитектурному паттерну MVVM (Model-View-ViewModel) "
    "с feature-first организацией кода. Корневое пространство имён — EduDev_Tracker "
    "(с символом подчёркивания; имя папки — «EduDev Tracker» с пробелом)."
))
add_body(doc, "Уровни архитектуры (от пользовательского интерфейса к данным):")
for layer in [
    "View (XAML-страницы) — отвечает за отображение. Содержится в Features/<Имя>/Views/. "
     "Страницы получают ViewModel через конструктор (DI) и устанавливают BindingContext.",
    "ViewModel — реализует логику представления. Унаследован от BaseViewModel. "
     "Использует [ObservableProperty] и [RelayCommand] из CommunityToolkit.Mvvm.",
    "Service (IXxxService) — содержит бизнес-правила и валидацию. Все сервисы зарегистрированы как Singleton.",
    "Repository : BaseRepository<T> — тонкий CRUD-слой. Инкапсулирует SQL-запросы к SQLite.",
    "DatabaseService — единый менеджер подключения SQLiteAsyncConnection. "
     "Создаёт все таблицы при старте, управляет миграциями.",
]:
    p = doc.add_paragraph()
    _apply_para_fmt(p, first_indent=False)
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f"— {layer}")
    _apply_font(run)

add_heading2(doc, "2.2 Внедрение зависимостей (DI)")
add_body(doc, (
    "Конфигурация выполняется в MauiProgram.cs в методе CreateMauiApp(). "
    "Применяется встроенный DI-контейнер .NET (Microsoft.Extensions.DependencyInjection)."
))
add_body(doc, "Время жизни компонентов:")
add_table(doc,
    headers=["Тип компонента", "Время жизни", "Примеры"],
    rows=[
        ["DatabaseService",        "Singleton", "Единственный экземпляр на весь процесс"],
        ["Репозитории (8 шт.)",    "Singleton", "HabitRepository, TaskRepository, NoteRepository и др."],
        ["Сервисы (12 шт.)",       "Singleton", "IHabitService, ITaskService, INotesService и др."],
        ["Seeders (2 шт.)",        "Singleton", "DemoDataSeeder, ActivitySeeder"],
        ["BackgroundTimerService", "Singleton", "Фоновый таймер Помодоро"],
        ["AudioManager",           "Singleton", "Plugin.Maui.Audio (AudioManager.Current)"],
        ["Pages (19 шт.)",         "Transient", "HabitsPage, TasksPage, NotesPage и др."],
        ["ViewModels (17 шт.)",    "Transient", "HabitsViewModel, TasksViewModel и др."],
    ],
    col_widths=[5.5, 3, 8],
)

add_heading2(doc, "2.3 База данных")
add_body(doc, (
    "Хранилище данных — SQLite 3.x через пакет sqlite-net-pcl. "
    "Используется единственное подключение SQLiteAsyncConnection, управляемое DatabaseService. "
    "Путь к файлу базы данных: {FileSystem.AppDataDirectory}/edudev-tracker.db3."
))
add_body(doc, (
    "База данных содержит 19 таблиц. Версионирование схемы осуществляется через PRAGMA user_version. "
    "Текущая версия схемы: 3 (константа Constants.CurrentSchemaVersion = 3). "
    "Для полнотекстового поиска по заметкам используется виртуальная таблица FTS5 (notes_fts) "
    "с тремя триггерами синхронизации (INSERT, UPDATE, DELETE)."
))
add_body(doc, "Состав таблиц базы данных:")
add_table(doc,
    headers=["Таблица", "Тип создания", "Назначение"],
    rows=[
        ["Profile",           "CreateTableAsync<T>()",  "Профили пользователей"],
        ["Tag",               "CreateTableAsync<T>()",  "Теги (общие для всех модулей)"],
        ["Habit",             "CreateTableAsync<T>()",  "Привычки"],
        ["HabitSchedule",     "CreateTableAsync<T>()",  "Расписание привычек (дни недели, время)"],
        ["habit_logs",        "Сырой SQL",              "Логи выполнения; UNIQUE(HabitId, LogDate)"],
        ["HabitTag",          "CreateTableAsync<T>()",  "M:N связь привычек и тегов"],
        ["Project",           "CreateTableAsync<T>()",  "Проекты (папки задач)"],
        ["TaskItem",          "CreateTableAsync<T>()",  "Задачи"],
        ["TaskRecurrence",    "CreateTableAsync<T>()",  "Правила повторения задач"],
        ["TaskTag",           "CreateTableAsync<T>()",  "M:N связь задач и тегов"],
        ["NoteCategory",      "CreateTableAsync<T>()",  "Категории заметок"],
        ["Note",              "CreateTableAsync<T>()",  "Заметки (поддержка Markdown)"],
        ["NoteAttachment",    "CreateTableAsync<T>()",  "Вложения к заметкам"],
        ["NoteTag",           "CreateTableAsync<T>()",  "M:N связь заметок и тегов"],
        ["NoteVersion",       "CreateTableAsync<T>()",  "История версий заметок"],
        ["PomodoroPreset",    "CreateTableAsync<T>()",  "Пресеты помодоро-таймера"],
        ["PomodoroSession",   "CreateTableAsync<T>()",  "Сессии помодоро"],
        ["ConversionHistory", "CreateTableAsync<T>() + индекс", "История конвертаций"],
        ["Reminder",          "CreateTableAsync<T>()",  "Напоминания к задачам"],
        ["notes_fts",         "Виртуальная FTS5",       "Полнотекстовый поиск по заметкам"],
    ],
    col_widths=[4.5, 5, 7],
)

add_heading2(doc, "2.4 Временные характеристики")
add_table(doc,
    headers=["Операция", "Характеристика"],
    rows=[
        ["Инициализация БД (первый запуск)", "200–400 мс (создание таблиц, FTS, индексов)"],
        ["Инициализация БД (повторный запуск)", "30–60 мс (таблицы уже существуют)"],
        ["Загрузка страницы (список привычек/задач)", "< 100 мс при < 500 записях"],
        ["Полнотекстовый поиск по заметкам (FTS5)", "< 50 мс, дебоунс 300 мс"],
        ["Тик помодоро-таймера (BackgroundTimerService)", "Каждые 1 с"],
        ["Звуковое предупреждение (Помодоро)", "За 60 с до конца фазы"],
        ["Дебоунс конвертеров (JSON/URL)", "400 мс после последнего ввода"],
        ["Дебоунс поиска по заметкам", "300 мс после последнего ввода"],
        ["Автосохранение заметки", "При каждом изменении содержимого"],
    ],
    col_widths=[9, 7.5],
)

add_heading2(doc, "2.5 Режимы работы")
add_table(doc,
    headers=["Режим", "Описание"],
    rows=[
        ["Однопользовательский (основной)", "Все данные привязаны к одному активному профилю на устройстве."],
        ["Офлайн (постоянный)",             "Приложение не обращается к серверу. Интернет нужен только для OAuth-входа."],
        ["Локальный (без аккаунта)",        "Анонимный профиль. Полная функциональность без регистрации."],
        ["DEBUG-режим",                     "Включает SQL-трассировку (SQLiteConnection.Tracer), расширенное логирование через Debug.WriteLine."],
    ],
    col_widths=[5.5, 11],
)

add_heading2(doc, "2.6 Механизмы контроля и самовосстановления")
for item in [
    ("Идемпотентность инициализации БД",
     "DatabaseService.InitAsync() содержит ранний выход (early return) при повторном вызове. "
     "Параллельный доступ защищён SemaphoreSlim(1,1)."),
    ("Безопасные миграции схемы",
     "Метод TryAddColumnAsync() перед ALTER TABLE проверяет наличие столбца через PRAGMA table_info(). "
     "Исключения при миграции перехватываются и логируются без прерывания запуска."),
    ("Нефатальное логирование",
     "Все блоки try-catch в сервисах логируют ошибку через Debug.WriteLine и продолжают "
     "работу. Только критические ошибки инициализации (DatabaseService) перенаправляют пользователя на экран входа."),
    ("Восстановление уведомлений",
     "При каждом запуске приложения вызывается INotificationService.RescheduleAllPendingAsync(), "
     "восстанавливающий все активные напоминания из БД."),
    ("Валидация входных данных",
     "Валидация происходит на двух уровнях: ViewModel (отображение ошибки пользователю до сохранения) "
     "и Service (выброс InvalidOperationException при нарушении бизнес-правил)."),
]:
    add_body(doc, f"{item[0]}. {item[1]}", space_before=2)

add_page_break(doc)

# ═══════════════════════════════════════════
# РАЗДЕЛ 3. ОБРАЩЕНИЕ К ПРИЛОЖЕНИЮ
# ═══════════════════════════════════════════

add_heading1(doc, "3 ОБРАЩЕНИЕ К ПРИЛОЖЕНИЮ")

add_heading2(doc, "3.1 Запуск приложения")
add_body(doc, (
    "CLI-параметры командной строки не используются. Конфигурация приложения "
    "выполняется исключительно через MauiProgram.cs (DI-регистрации) и MAUI Preferences "
    "(настройки пользователя на устройстве)."
))
add_body(doc, "Способы запуска:")
add_table(doc,
    headers=["Метод", "Команда / Действие"],
    rows=[
        ["Из среды разработки (Windows)",   "dotnet build \"EduDev Tracker/EduDev Tracker.csproj\" -t:Run -f net10.0-windows10.0.19041.0"],
        ["Сборка без запуска (Windows)",    "dotnet build \"EduDev Tracker/EduDev Tracker.csproj\" -f net10.0-windows10.0.19041.0"],
        ["Сборка для Android",              "dotnet build \"EduDev Tracker/EduDev Tracker.csproj\" -f net10.0-android"],
        ["Публикация (Windows)",            "dotnet publish -f net10.0-windows10.0.19041.0 -c Release"],
        ["Исполняемый файл (после публикации)", "Запуск EduDev Tracker.exe из папки публикации"],
    ],
    col_widths=[5, 11.5],
)

add_heading2(doc, "3.2 Последовательность инициализации")
add_body(doc, (
    "Точка входа — App.xaml.cs, метод InitializeDatabaseAsync(). "
    "Вызывается асинхронно из конструктора App. Последовательность шагов:"
))
steps = [
    ("Шаг 1", "DatabaseService.InitAsync()",
     "Открытие файла edudev-tracker.db3 (создание при первом запуске). "
     "Создание 19 таблиц, FTS5-индекса и триггеров. "
     "Проверка PRAGMA user_version, применение отсутствующих миграций."),
    ("Шаг 2", "SessionService.GetProfileId()",
     "Чтение ProfileId из MAUI Preferences. "
     "Если значение равно 0 — переход к шагу AuthPage (пользователь не авторизован)."),
    ("Шаг 3", "ProfileRepository.GetByIdAsync(profileId)",
     "Проверка существования профиля в БД. "
     "Если профиль не найден — очистка Preferences, переход к AuthPage."),
    ("Шаг 4", "INotificationService.RescheduleAllPendingAsync(profileId)",
     "Восстановление всех запланированных уведомлений о привычках и задачах. "
     "Ошибка логируется как [App Init] Reschedule failed, не прерывает запуск."),
    ("Шаг 5", "DemoDataSeeder.SeedIfNeededAsync(profileId)",
     "Однократное заполнение справочных данных (чит-листы). "
     "Повторный вызов не выполняет сид (идемпотентно)."),
    ("Шаг 6", "ActivitySeeder.SeedIfNeededAsync(profileId)",
     "Заполнение демонстрационных данных активности за текущий месяц. "
     "Может быть отключено в коде."),
    ("Шаг 7", "MainPage = new AppShell()",
     "Установка основного UI (flyout-навигация с 8 разделами). "
     "При ошибке на любом шаге вместо этого вызывается GoToAuth() → NavigationPage(AuthPage)."),
]
for s, method, desc in steps:
    p = doc.add_paragraph()
    _apply_para_fmt(p, first_indent=False)
    p.paragraph_format.left_indent = Cm(1.25)
    r1 = p.add_run(f"{s}: {method} — ")
    _apply_font(r1, bold=True)
    r2 = p.add_run(desc)
    _apply_font(r2)

add_heading2(doc, "3.3 Конфигурация DI (MauiProgram.cs)")
add_body(doc, (
    "Все компоненты регистрируются в MauiProgram.cs методом CreateMauiApp(). "
    "Порядок регистрации: сначала инфраструктура (DatabaseService), "
    "затем репозитории, seeders, сервисы, фоновые сервисы, страницы и ViewModels."
))

add_heading2(doc, "3.4 Навигация")
add_body(doc, (
    "Навигация реализована через Shell-based routing (.NET MAUI Shell). "
    "ViewModels не обращаются к Shell.Current напрямую — используют INavigationService."
))
add_table(doc,
    headers=["Метод INavigationService", "Описание"],
    rows=[
        ["GoToAsync(route)",              "Переход по зарегистрированному маршруту"],
        ["GoToAsync(route, parameters)",  "Переход с передачей параметров (QueryProperty)"],
        ["GoBackAsync()",                 "Возврат на предыдущую страницу (Shell.GoToAsync(\"..\"))"],
        ["PushModalAsync(page)",          "Открытие страницы как модального окна"],
        ["PopModalAsync()",               "Закрытие верхнего модального окна"],
        ["SwitchToModuleAsync(flyoutRoute)", "Переключение раздела flyout-меню (Dashboard, Habits и др.)"],
    ],
    col_widths=[6.5, 10],
)
add_body(doc, "Зарегистрированные маршруты (AppShell.xaml.cs):")
add_table(doc,
    headers=["Маршрут", "Страница назначения"],
    rows=[
        ["HabitDetailsPage",    "Детальный просмотр привычки (параметр: habitId)"],
        ["ArchivedHabitsPage",  "Архив привычек"],
        ["CreateHabitPage",     "Форма создания привычки"],
        ["TaskDetailsPage",     "Детальный просмотр задачи (параметр: taskId)"],
        ["AddTaskPage",         "Форма создания задачи"],
        ["ReminderSettingsPage","Настройка напоминания для задачи (параметр: taskId)"],
    ],
    col_widths=[5.5, 11],
)

add_heading2(doc, "3.5 Расположение файла базы данных")
add_table(doc,
    headers=["Платформа", "Путь к файлу"],
    rows=[
        ["Windows", "%LOCALAPPDATA%\\Packages\\com.companyname.edudevtracker\\LocalState\\edudev-tracker.db3"],
        ["Android", "/data/data/com.companyname.edudevtracker/files/edudev-tracker.db3"],
        ["iOS",     "{NSDocumentDirectory}/edudev-tracker.db3 (Sandbox приложения)"],
        ["macOS",   "~/Library/Containers/com.companyname.edudevtracker/Data/.../edudev-tracker.db3"],
    ],
    col_widths=[3, 13.5],
)

add_page_break(doc)

# ═══════════════════════════════════════════
# РАЗДЕЛ 4. ВХОДНЫЕ И ВЫХОДНЫЕ ДАННЫЕ
# ═══════════════════════════════════════════

add_heading1(doc, "4 ВХОДНЫЕ И ВЫХОДНЫЕ ДАННЫЕ")

add_heading2(doc, "4.1 Входные данные")
add_body(doc, "Приложение получает входные данные из следующих источников:")
add_table(doc,
    headers=["Источник", "Тип данных", "Описание"],
    rows=[
        ["UI (XAML-контролы)",   "Строки, числа, булевы, DateTime", "Ввод через Editor, Entry, Picker, DatePicker, TimePicker. Биндинг через [ObservableProperty]."],
        ["SQLite БД",            "Объекты C# (ORM-маппинг)",        "Чтение записей при загрузке страниц. GetAllAsync(), GetActiveAsync(profileId), GetWithChildrenAsync()."],
        ["MAUI Preferences",     "int (ProfileId), bool, string",   "Хранит текущий ProfileId и настройки пользователя (язык, уведомления, звук)."],
        ["Файловая система",     "Байтовые потоки файлов",          "Вложения к заметкам. Файлы копируются в AppDataDirectory через FilePicker."],
        ["Plugin.LocalNotification", "NotificationTappedEventArgs", "Событие нажатия на уведомление возвращает параметры маршрута навигации."],
    ],
    col_widths=[4, 4, 8.5],
)

add_heading2(doc, "4.2 Выходные данные")
add_table(doc,
    headers=["Получатель", "Формат", "Описание"],
    rows=[
        ["SQLite БД",        "SQL INSERT/UPDATE/DELETE",          "Запись через SaveAsync(entity) / SaveWithChildrenAsync(). Метод определяет INSERT (Id==0) или UPDATE (Id>0) через рефлексию."],
        ["MAUI Preferences", "Key-Value пары",                   "Сохранение ProfileId (ключ: profile_id), настроек (ключи: setting_*_{profileId})."],
        ["Push-уведомления", "Plugin.LocalNotification API",     "Планирование уведомлений с датой, заголовком, телом и параметрами навигации."],
        ["Файл .xlsx",       "ClosedXML (Excel Open XML)",       "Экспорт аналитики через IExportService.ExportToExcelAsync(). Файл сохраняется в кэш."],
        ["Файл .md / .txt",  "UTF-8 текст",                     "Экспорт содержимого заметки через системный диалог шеринга."],
        ["PDF",              "Системный диалог печати/шеринга",  "Генерируется через платформенный Share API без промежуточного файла."],
        ["Аудио",            "Plugin.Maui.Audio",                "Воспроизведение .mp3/.wav файлов из Resources/Raw (звуки помодоро)."],
        ["Экран (UI)",       "MAUI XAML Binding",                "Обновление интерфейса через INotifyPropertyChanged. Данные из ViewModel передаются в View."],
    ],
    col_widths=[3.5, 4, 9],
)

add_heading2(doc, "4.3 Схема данных — ключевые сущности")
add_body(doc, "4.3.1 Сущность Profile (профиль пользователя):")
add_table(doc,
    headers=["Поле", "Тип SQLite", "Описание"],
    rows=[
        ["Id",           "INTEGER PK",  "Идентификатор профиля"],
        ["Name",         "TEXT",        "Отображаемое имя"],
        ["Email",        "TEXT NULL",   "Email (null в локальном режиме)"],
        ["PasswordHash", "TEXT NULL",   "BCrypt-хэш пароля"],
        ["IsLocalMode",  "INTEGER",     "1 = анонимный локальный профиль"],
        ["OAuthProvider","TEXT NULL",   "«google» или «github» (после OAuth-входа)"],
        ["OAuthId",      "TEXT NULL",   "Идентификатор OAuth-провайдера"],
        ["CreatedAt",    "TEXT",        "Дата создания (ISO-8601)"],
        ["LastLoginAt",  "TEXT",        "Последний вход"],
    ],
    col_widths=[3.5, 3, 10],
)
add_body(doc, "4.3.2 Сущность Habit (привычка):")
add_table(doc,
    headers=["Поле", "Тип SQLite", "Описание"],
    rows=[
        ["Id",          "INTEGER PK",  "Идентификатор"],
        ["ProfileId",   "INTEGER FK",  "Ссылка на Profile.Id"],
        ["Title",       "TEXT",        "Название (до 200 символов)"],
        ["Type",        "INTEGER",     "Перечисление: 0=Binary, 1=Quantitative, 2=Time"],
        ["TargetValue", "REAL NULL",   "Целевое значение (для Quantitative/Time)"],
        ["TargetUnit",  "TEXT NULL",   "Единица измерения"],
        ["Icon",        "TEXT NULL",   "Имя файла иконки (default_icon.png и др.)"],
        ["IsArchived",  "INTEGER",     "1 = архивирована"],
        ["IsFrozen",    "INTEGER",     "1 = заморожена"],
        ["CreatedAt",   "TEXT",        "Дата создания"],
        ["UpdatedAt",   "TEXT",        "Дата последнего изменения"],
    ],
    col_widths=[3.5, 3, 10],
)
add_body(doc, "4.3.3 Сущность TaskItem (задача):")
add_table(doc,
    headers=["Поле", "Тип SQLite", "Описание"],
    rows=[
        ["Id",         "INTEGER PK",  "Идентификатор"],
        ["ProfileId",  "INTEGER FK",  "Ссылка на Profile.Id"],
        ["Title",      "TEXT",        "Название (до 300 символов)"],
        ["Priority",   "INTEGER",     "0=Low, 1=Medium, 2=High, 3=Urgent"],
        ["Status",     "INTEGER",     "0=Open, 1=InProgress, 2=Done, 3=Cancelled"],
        ["DueAt",      "TEXT NULL",   "Дедлайн (ISO-8601)"],
        ["ReminderAt", "INTEGER NULL","Время напоминания (Unix timestamp)"],
        ["IsArchived", "INTEGER",     "1 = архивирована"],
        ["CreatedAt",  "TEXT",        "Дата создания"],
    ],
    col_widths=[3.5, 3, 10],
)
add_body(doc, "4.3.4 Кодирование специальных значений:")
add_table(doc,
    headers=["Значение", "Кодирование", "Пример"],
    rows=[
        ["Пароль",        "BCrypt (cost 11)",           "«$2a$11$xyz...» (60 символов)"],
        ["Дни недели",    "Битовая маска (int)",         "Пн=1, Вт=2, Ср=4, Чт=8, Пт=16, Сб=32, Вс=64; все 7 дней = 127"],
        ["Дата/время",    "ISO-8601 строка или unix int","«2026-06-22T10:30:00» / 1750589400"],
        ["ProfileId",     "int в MAUI Preferences",     "Preferences.Set(\"profile_id\", 42)"],
        ["Настройки",     "Prefixed key string",        "Ключ: «setting_smart_reminders_42» (42 = profileId)"],
    ],
    col_widths=[3.5, 4.5, 8.5],
)

add_page_break(doc)

# ═══════════════════════════════════════════
# РАЗДЕЛ 5. СООБЩЕНИЯ
# ═══════════════════════════════════════════

add_heading1(doc, "5 СООБЩЕНИЯ")

add_heading2(doc, "5.1 Системные сообщения инициализации")
add_body(doc, (
    "Сообщения выводятся через System.Diagnostics.Debug.WriteLine() "
    "и доступны в окне «Output» Visual Studio при запуске в DEBUG-конфигурации."
))
add_table(doc,
    headers=["Сообщение", "Источник", "Действие программиста"],
    rows=[
        ["[App Init] profileId из Preferences = {id}",
         "App.xaml.cs",
         "Информационное. Показывает загруженный ProfileId."],
        ["[App Init] profile из БД = {name|null}",
         "App.xaml.cs",
         "Если «null» — профиль не найден в БД, пользователь будет перенаправлен на AuthPage."],
        ["[App Init] Reschedule failed: {msg}",
         "App.xaml.cs",
         "Ошибка восстановления уведомлений. Не критично; проверить разрешения уведомлений."],
        ["[App Init] Seed failed: {msg}",
         "App.xaml.cs",
         "Ошибка сида данных. Проверить DemoDataSeeder.SeedIfNeededAsync()."],
        ["[App Init] EXCEPTION: {ex}",
         "App.xaml.cs",
         "Критическая ошибка инициализации. Проверить стек трассировки, скорее всего ошибка в DatabaseService."],
    ],
    col_widths=[6, 3, 7.5],
)

add_heading2(doc, "5.2 Сообщения базы данных")
add_table(doc,
    headers=["Сообщение", "Источник", "Действие программиста"],
    rows=[
        ["DB: opened",          "DatabaseService", "Подключение к БД установлено успешно."],
        ["DB: {TableName}",     "DatabaseService", "Таблица создана (или уже существовала). Норма при каждом запуске."],
        ["DB: tables done",     "DatabaseService", "Все таблицы созданы."],
        ["DB: fts done",        "DatabaseService", "FTS5-индекс и триггеры созданы."],
        ["DB: migrate done",    "DatabaseService", "Миграции схемы применены."],
        ["[SQL] {query}",       "DatabaseService (Tracer)", "Полный SQL-запрос. Только в DEBUG. Для диагностики N+1 запросов."],
        ["[Migrate] {table}.{col}: {msg}", "DatabaseService", "Ошибка добавления столбца. Столбец мог уже существовать — не критично."],
    ],
    col_widths=[5.5, 3.5, 7.5],
)

add_heading2(doc, "5.3 Сообщения модулей (Debug)")
add_table(doc,
    headers=["Сообщение", "Модуль", "Действие программиста"],
    rows=[
        ["[HabitsViewModel] Error: {msg}",      "Привычки",   "Ошибка загрузки списка привычек. Проверить HabitService.GetActiveAsync()."],
        ["[Archive] Ошибка: {msg}",              "Привычки",   "Ошибка загрузки архива. Проверить HabitService.GetArchivedAsync()."],
        ["[GetActiveAsync] CRASH: {ex}",         "Задачи",     "Падение при загрузке задач. Проверить TaskRepository.GetActiveAsync()."],
        ["[TasksViewModel.Load] CRASH: {ex}",    "Задачи",     "Падение в TasksViewModel. Проверить LoadAsync() и FromModel()."],
        ["[FromModel ERROR] TaskId={id}",        "Задачи",     "Ошибка маппинга TaskItem → TaskItemViewModel. Проверить поля TaskItem."],
        ["[ProcessAsync] ERROR: {ex}",           "Задачи",     "Ошибка в RecurrenceProcessorService. Проверить CalcNextDue()."],
        ["[SaveNote] {ex}",                      "Заметки",    "Ошибка сохранения заметки. Проверить NotesService.SaveAsync()."],
        ["[VERSION] SKIP: контент не изменился", "Заметки",    "Версия не сохранена (содержимое то же). Норма."],
        ["[AddAttachment] {ex}",                 "Заметки",    "Ошибка прикрепления файла. Проверить разрешения на хранилище."],
        ["[OpenAttachment] {ex}",                "Заметки",    "Файл не открылся. Проверить наличие файла и приложения-обработчика."],
        ["[LoadMarkedJs] {ex}",                  "Заметки",    "Ошибка загрузки JS-парсера Markdown. Проверить Resources/Raw/marked.min.js."],
        ["[OnTimerTick] {ex}",                   "Помодоро",   "Ошибка тика таймера. Проверить BackgroundTimerService."],
        ["[AudioService] Failed to play {sound}","Помодоро",   "Аудиофайл не воспроизведён. Проверить Resources/Raw/{sound}.mp3."],
        ["[Analytics] {ex}",                     "Аналитика",  "Ошибка загрузки аналитики. Проверить IAnalyticsService."],
        ["[Nav] Модуль '{route}' не найден",     "Навигация",  "Маршрут flyout не найден. Проверить AppShell.xaml (FlyoutItem Route)."],
        ["[OAuth Google] {ex}",                  "Авторизация","Ошибка Google OAuth. Проверить OAuthService и конфигурацию OAuth-клиента."],
        ["[OAuth GitHub] {ex}",                  "Авторизация","Ошибка GitHub OAuth. Аналогично."],
        ["[Dashboard] {ex}",                     "Дашборд",    "Ошибка загрузки дашборда. Проверить DashboardViewModel."],
    ],
    col_widths=[6.5, 2.5, 7.5],
)

add_heading2(doc, "5.4 Исключения бизнес-логики")
add_table(doc,
    headers=["Тип исключения", "Сообщение / код", "Источник", "Обработка"],
    rows=[
        ["InvalidOperationException", "\"Call InitAsync() first\"",     "DatabaseService.Connection", "Вызван до InitAsync(). Переупорядочить запуск."],
        ["InvalidOperationException", "\"Название привычки обязательно!\"", "HabitService.CreateAsync()", "Пустое название. Валидация в ViewModel."],
        ["InvalidOperationException", "\"EMAIL_TAKEN\"",                "AuthService.SignUpAsync()",  "Email уже зарегистрирован. ViewModel показывает ошибку."],
        ["InvalidOperationException", "\"USER_NOT_FOUND\"",             "AuthService.LoginAsync()",   "Email не найден. ViewModel показывает ошибку."],
        ["InvalidOperationException", "\"WRONG_PASSWORD\"",             "AuthService.LoginAsync()",   "Неверный пароль. ViewModel показывает ошибку."],
        ["NotImplementedException",   "(ConvertBack не реализован)",    "XAML-конверторы",            "ConvertBack() в IValueConverter не используется. Не исправлять."],
    ],
    col_widths=[4, 4.5, 3.5, 4.5],
)

add_heading2(doc, "5.5 Пользовательские диалоги (DisplayAlertAsync)")
add_body(doc, (
    "Диалоги отображаются через Shell.Current.DisplayAlertAsync(). "
    "Используются для подтверждения деструктивных операций и информирования об ошибках."
))
add_table(doc,
    headers=["Заголовок диалога", "Текст", "Кнопки", "Контекст вызова"],
    rows=[
        ["«Очистить историю»",        "«Удалить все конверсии за сегодня?»",                          "«Очистить» / «Отмена»",  "ConvertersViewModel — очистка истории за сегодня"],
        ["«Очистить всю историю»",    "«Удалить всю историю конверсий без возможности восстановления?»", "«Удалить всё» / «Отмена»","ConvertersViewModel — полная очистка истории"],
        ["«Удалить заметку»",         "«Удалить «{title}»?»",                                          "«Удалить» / «Отмена»",   "NotesViewModel — удаление заметки"],
        ["«Удалить вложение»",        "«Удалить файл «{filename}»?»",                                  "«Удалить» / «Отмена»",   "NotesViewModel — удаление вложения"],
        ["(без заголовка)",           "«Не удалось открыть файл»",                                    "«OK»",                   "NotesViewModel — ошибка открытия файла"],
        ["«Подтвердите свои действия»","(текст удаления/архивации привычки)",                          "«Да» / «Нет»",           "HabitDetailsViewModel, ArchivedHabitsViewModel"],
        ["«Ошибка»",                  "«Некорректный ввод»",                                          "«Принято»",              "HabitsViewModel — ввод нечислового значения"],
        ["«Ошибка»",                  "«Название задачи обязательно»",                                "«ОК»",                   "AddTaskViewModel — пустое название"],
        ["«Ошибка»",                  "«Название задачи должно быть больше 3 символов»",              "«ОК»",                   "AddTaskViewModel — слишком короткое название"],
        ["«Ошибка»",                  "«Дедлайн не может существовать в прошлом»",                    "«ОК»",                   "AddTaskViewModel — прошедшая дата"],
        ["«Ошибка»",                  "«Выберите хотя бы один день недели»",                          "«ОК»",                   "AddTaskViewModel — еженедельное повторение без дней"],
        ["«Успех!»",                  "«Задача создана!»",                                            "«ОК»",                   "AddTaskViewModel — задача успешно создана"],
        ["«Подтвердите свои действия»","(текст удаления/архивации задачи)",                           "«Да» / «Нет»",           "TasksViewModel, TaskDetailsViewModel"],
        ["«Ошибка загрузки»",         "{ex.Message}",                                                 "«OK»",                   "TasksViewModel — ошибка загрузки с показом текста исключения"],
    ],
    col_widths=[3.5, 5, 2.5, 5.5],
)

add_page_break(doc)

# ═══════════════════════════════════════════
# ПРИЛОЖЕНИЕ: Диаграмма архитектуры
# ═══════════════════════════════════════════

add_heading1(doc, "ПРИЛОЖЕНИЕ А. ДИАГРАММА АРХИТЕКТУРЫ ПРИЛОЖЕНИЯ")
add_body(doc, "Общая схема компонентов и их взаимодействия:")

arch_lines = [
    "Запуск приложения",
    "    └── App.xaml.cs → InitializeDatabaseAsync()",
    "            ├── DatabaseService.InitAsync()",
    "            │       ├── Открытие edudev-tracker.db3",
    "            │       ├── Создание 19 таблиц + FTS5",
    "            │       └── Миграции (PRAGMA user_version)",
    "            ├── SessionService → MAUI Preferences",
    "            └── AppShell (8 flyout-разделов)",
    "",
    "AppShell",
    "    ├── DashboardPage      → DashboardViewModel  → IHabitService, ITaskService",
    "    ├── HabitsPage         → HabitsViewModel     → IHabitService    → HabitRepository",
    "    ├── TasksPage          → TasksViewModel      → ITaskService     → TaskRepository",
    "    ├── NotesPage          → NotesViewModel      → INotesService    → NoteRepository (FTS5)",
    "    ├── PomodoroPage       → PomodoroViewModel   → IPomodoroService → PomodoroRepository",
    "    │                                            └── BackgroundTimerService",
    "    ├── ConvertersPage     → ConvertersViewModel → IConversionService → ConversionRepository",
    "    ├── AnalyticsPage      → AnalyticsViewModel  → IAnalyticsService",
    "    └── ProfilePage        → ProfileViewModel    → IAuthService, SessionService",
    "",
    "Все Repository → BaseRepository<T> → DatabaseService → SQLiteAsyncConnection",
    "Все ViewModel  → BaseViewModel (IsBusy, Title, InitializeAsync)",
    "Все Services   → Singleton | Pages/ViewModels → Transient",
]

for line in arch_lines:
    p = doc.add_paragraph()
    _apply_para_fmt(p, first_indent=False)
    run = p.add_run(line)
    run.font.name  = "Courier New"
    run.font.size  = Pt(11)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

doc.add_paragraph()

add_heading1(doc, "ПРИЛОЖЕНИЕ Б. ИСТОРИЯ МИГРАЦИЙ СХЕМЫ БАЗЫ ДАННЫХ")
add_table(doc,
    headers=["Версия", "Изменения", "Метод"],
    rows=[
        ["0 → 1", "Начальная схема (все таблицы)", "CreateTableAsync<T>() для каждой сущности"],
        ["1 → 2", "Добавлены поля Profile.OAuthProvider (TEXT) и Profile.OAuthId (TEXT)", "TryAddColumnAsync('profiles', 'OAuthProvider', 'TEXT NULL')"],
        ["2 → 3", "Добавлено поле TaskItem.ReminderAt (INTEGER, unix timestamp)",         "TryAddColumnAsync('tasks', 'ReminderAt', 'INTEGER NULL')"],
        ["3 → ?", "Не определено", "При необходимости: инкрементировать Constants.CurrentSchemaVersion и добавить ветку в MigrateAsync()"],
    ],
    col_widths=[2, 8, 6.5],
)

# ─────────────────────────────────────────────
# Сохранение
# ─────────────────────────────────────────────
out_path = r"d:\C# repositories\EduDev Tracker\docs\ProgrammerManual.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
